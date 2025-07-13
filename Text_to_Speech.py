import os
import time
from docx import Document
from io import BytesIO
import pyttsx3  # Offline fallback
import requests  # For internet check
import threading

# Configuration
DOCX_PATH = input("Enter the path to your DOCX file: ")  # Path to DOCX file
if not os.path.exists(DOCX_PATH):
    print("File not found. Please check the path and try again.")
    exit(1)
SPEECH_DELAY = 0.08  # Character print delay
ONLINE_FIRST = True  # Try online first, fallback to offline

def check_internet():
    """Check internet connection"""
    try:
        requests.get("http://www.google.com", timeout=3)
        return True
    except:
        return False

def extract_docx_text(file_path):
    """Extract text from DOCX with paragraphs and tables"""
    doc = Document(file_path)
    full_text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    for table in doc.tables:
        full_text.append("\n[Table Content]")
        for row in table.rows:
            full_text.append(" | ".join(cell.text for cell in row.cells))
    
    return "\n".join(full_text)

def offline_tts_speak(text):
    """Reliable offline TTS using pyttsx3"""
    try:
        engine = pyttsx3.init()
        # Configure voice
        engine.setProperty('rate', 150)  # Speaking speed
        engine.setProperty('volume', 0.9)  # Volume (0-1)
        voices = engine.getProperty('voices')
        engine.setProperty('voice', voices[0].id)  # Change voice if needed
        
        # Print and speak simultaneously
        engine.say(text)
        
        # Start speech in background thread
        speech_thread = threading.Thread(target=engine.runAndWait)
        speech_thread.start()
        
        # Print text with delay
        for char in text:
            print(char, end='', flush=True)
            time.sleep(SPEECH_DELAY)
        print()
        
        speech_thread.join()
        return True
        
    except Exception as e:
        print(f"\n[Offline TTS Error] {e}")
        return False

def online_tts_speak(text):
    """Try online gTTS with fallback"""
    try:
        from gtts import gTTS
        temp_file = "temp_speech.mp3"
        
        # Generate speech
        tts = gTTS(text=text, lang='en')
        tts.save(temp_file)
        
        # Print with delay while playing
        def play_audio():
            import playsound
            playsound.playsound(temp_file)
            os.remove(temp_file)
        
        # Start audio thread
        audio_thread = threading.Thread(target=play_audio)
        audio_thread.start()
        
        # Print text
        for char in text:
            print(char, end='', flush=True)
            time.sleep(SPEECH_DELAY)
        print()
        
        audio_thread.join()
        return True
        
    except Exception as e:
        print(f"\n[Online TTS Error] {e}")
        return False

def main():
    # Extract text
    try:
        text_content = extract_docx_text(DOCX_PATH)
        if not text_content.strip():
            print("Document is empty!")
            return
    except Exception as e:
        print(f"Error reading document: {e}")
        return
    
    # Select TTS method
    if ONLINE_FIRST and check_internet():
        print("Using online TTS...")
        if not online_tts_speak(text_content):
            print("Falling back to offline TTS...")
            offline_tts_speak(text_content)
    else:
        print("Using offline TTS...")
        offline_tts_speak(text_content)

if __name__ == "__main__":
    main()